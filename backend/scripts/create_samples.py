from openpyxl import Workbook
from openpyxl.styles import PatternFill
from docx import Document
import os

def create_samples():
    # 1. Create Sample Questionnaire
    wb = Workbook()
    ws = wb.active
    ws.title = "Sample Safety Form"
    ws["A1"] = "Question"
    ws["B1"] = "Answer"
    
    questions = [
        "What is the company policy on PPE?",
        "Who is the designated Site Safety Manager?",
        "Describe the fire prevention protocol."
    ]
    
    yellow = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    
    for i, q in enumerate(questions, start=2):
        ws[f"A{i}"] = q
        ws[f"B{i}"].fill = yellow
        
    wb.save("samples/sample_questionnaire.xlsx")
    print("✅ Created sample_questionnaire.xlsx")
    
    # 2. Create Sample Knowledge Doc
    doc = Document()
    doc.add_heading('Acme Corp Safety Manual', 0)
    
    doc.add_heading('1. Personal Protective Equipment (PPE)', level=1)
    doc.add_paragraph(
        'All employees must wear hard hats, high-visibility vests, and steel-toed boots at all times. '
        'Safety glasses are required when operating machinery.'
    )
    
    doc.add_heading('2. Site Safety Management', level=1)
    doc.add_paragraph(
        'The designated Site Safety Manager (SSM) is responsible for daily inspections and '
        'toolbox talks. The current SSM is John Doe.'
    )
    
    doc.add_heading('3. Fire Prevention', level=1)
    doc.add_paragraph(
        'Fire extinguishers must be placed within 75 feet of every working area. '
        'Hot work permits are required for all welding operations.'
    )
    
    doc.save("samples/sample_knowledge.docx")
    print("✅ Created sample_knowledge.docx")

if __name__ == "__main__":
    create_samples()
